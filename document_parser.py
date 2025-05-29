import fitz  # PyMuPDF for PDF
import docx  # python-docx for DOCX
import re
from typing import Optional, Dict, Any
from pathlib import Path

class DocumentParser:
    def __init__(self):
        """Initialize document parser supporting multiple formats"""
        self.supported_formats = {'.pdf', '.doc', '.docx', '.txt'}
        self.groq_processor = None
        try:
            from groq_processor import GroqProcessor
            self.groq_processor = GroqProcessor()
        except Exception:
            pass
    
    def extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from various document formats
        
        Args:
            file_content: File content as bytes
            filename: Original filename to determine format
            
        Returns:
            Extracted text as string
        """
        file_ext = Path(filename).suffix.lower()
        
        if file_ext == '.pdf':
            return self._extract_from_pdf(file_content)
        elif file_ext == '.docx':
            return self._extract_from_docx(file_content)
        elif file_ext == '.doc':
            return self._extract_from_doc(file_content)
        elif file_ext == '.txt':
            return self._extract_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_from_pdf(self, pdf_content: bytes) -> str:
        """Extract text from PDF using PyMuPDF"""
        try:
            pdf_document = fitz.open(stream=pdf_content, filetype="pdf")
            extracted_text = ""
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                page_text = page.get_text()
                extracted_text += self._clean_text(page_text) + "\n"
            
            pdf_document.close()
            
            # Enhance with Groq if available
            if self.groq_processor and extracted_text.strip():
                try:
                    return self.groq_processor.enhance_resume_text(extracted_text)
                except Exception:
                    pass
            
            return extracted_text.strip()
            
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_from_docx(self, docx_content: bytes) -> str:
        """Extract text from DOCX using python-docx"""
        try:
            import io
            from docx import Document
            
            # Create file-like object from bytes
            docx_file = io.BytesIO(docx_content)
            doc = Document(docx_file)
            
            extracted_text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + " "
                    extracted_text += "\n"
            
            return self._clean_text(extracted_text)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_from_doc(self, doc_content: bytes) -> str:
        """Extract text from DOC files (legacy Word format)"""
        try:
            # For DOC files, we need python-docx2txt or similar
            # This is a basic implementation that attempts text extraction
            import io
            
            # Try to decode as text (this is a fallback method)
            # For proper DOC support, you would need additional libraries like python-docx2txt
            try:
                text = doc_content.decode('utf-8', errors='ignore')
                return self._clean_text(text)
            except:
                # Alternative: Extract readable text portions
                readable_text = ""
                for byte in doc_content:
                    if 32 <= byte <= 126:  # Printable ASCII range
                        readable_text += chr(byte)
                    else:
                        readable_text += " "
                
                return self._clean_text(readable_text)
                
        except Exception as e:
            raise Exception(f"Failed to extract text from DOC: {str(e)}")
    
    def _extract_from_txt(self, txt_content: bytes) -> str:
        """Extract text from TXT files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = txt_content.decode(encoding)
                    return self._clean_text(text)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text = txt_content.decode('utf-8', errors='replace')
            return self._clean_text(text)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from TXT: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace and newlines
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere with processing
        text = re.sub(r'[^\w\s\.\,\-\(\)\@\+\/\&\%\$\#\!\?\:\;]', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        
        # Remove leading/trailing whitespace
        return text.strip()
    
    def validate_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Validate document and return validation results
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Dictionary containing validation results
        """
        file_ext = Path(filename).suffix.lower()
        
        validation_result = {
            'is_valid': False,
            'file_format': file_ext,
            'file_size': len(file_content),
            'messages': [],
            'can_extract_text': False
        }
        
        # Check if format is supported
        if file_ext not in self.supported_formats:
            validation_result['messages'].append(f"Unsupported file format: {file_ext}")
            return validation_result
        
        # Check file size (limit to 10MB)
        max_size = 10 * 1024 * 1024  # 10MB
        if len(file_content) > max_size:
            validation_result['messages'].append(f"File too large: {len(file_content)} bytes (max: {max_size})")
            return validation_result
        
        # Check if file is empty
        if len(file_content) == 0:
            validation_result['messages'].append("File is empty")
            return validation_result
        
        # Try to extract text to validate content
        try:
            extracted_text = self.extract_text_from_file(file_content, filename)
            if extracted_text.strip():
                validation_result['can_extract_text'] = True
                validation_result['is_valid'] = True
                validation_result['messages'].append("Document validation successful")
            else:
                validation_result['messages'].append("Warning: No readable text found in document")
                validation_result['is_valid'] = True  # Still valid, just no text
                
        except Exception as e:
            validation_result['messages'].append(f"Text extraction failed: {str(e)}")
        
        return validation_result
    
    def get_document_metadata(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract metadata from document
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Dictionary containing document metadata
        """
        file_ext = Path(filename).suffix.lower()
        
        metadata = {
            'filename': filename,
            'file_format': file_ext,
            'file_size': len(file_content),
            'page_count': 0,
            'word_count': 0,
            'character_count': 0
        }
        
        try:
            # Extract text to get basic statistics
            text = self.extract_text_from_file(file_content, filename)
            metadata['character_count'] = len(text)
            metadata['word_count'] = len(text.split()) if text else 0
            
            # Format-specific metadata
            if file_ext == '.pdf':
                try:
                    pdf_document = fitz.open(stream=file_content, filetype="pdf")
                    metadata['page_count'] = pdf_document.page_count
                    
                    # Get PDF metadata if available
                    pdf_meta = pdf_document.metadata
                    if pdf_meta:
                        metadata.update({
                            'title': pdf_meta.get('title', ''),
                            'author': pdf_meta.get('author', ''),
                            'creator': pdf_meta.get('creator', ''),
                            'creation_date': pdf_meta.get('creationDate', '')
                        })
                    
                    pdf_document.close()
                except Exception:
                    pass
            
            elif file_ext == '.docx':
                try:
                    import io
                    from docx import Document
                    docx_file = io.BytesIO(file_content)
                    doc = Document(docx_file)
                    
                    metadata['page_count'] = 1  # DOCX doesn't have fixed pages
                    metadata['paragraph_count'] = len(doc.paragraphs)
                    metadata['table_count'] = len(doc.tables)
                    
                    # Get core properties if available
                    core_props = doc.core_properties
                    metadata.update({
                        'title': core_props.title or '',
                        'author': core_props.author or '',
                        'created': str(core_props.created) if core_props.created else '',
                        'modified': str(core_props.modified) if core_props.modified else ''
                    })
                    
                except Exception:
                    pass
            
        except Exception as e:
            metadata['error'] = str(e)
        
        return metadata